"""Unit tests for stabler.api._bid_package (WP-306).

Assembly is Frappe-free and always runs. The docx render test skips when
python-docx is not installed (prod bench must add it — see go-live checklist).

    cd /path/to/stabler && PYTHONPATH=$PWD python3 -m unittest stabler.tests.test_bid_package -v
"""

from __future__ import annotations

import os
import tempfile
import unittest
from datetime import date

from stabler.api._bid_package import assemble_bid_package, fmt_date, fmt_money

_INTAKE = {"lot_no": "26111006497638", "buyer": "Maktab #1", "bid_deadline": "2026-07-08 11:47:09"}
_PNL = {
	"bid_price": 6300000000.0,
	"vat": 900000000.0,
	"net_revenue": 5400000000.0,
	"landed_goods": 4000000000.0,
	"profit": 800000000.0,
	"ostatok": 500000000.0,
	"margin_on_revenue_pct": 14.8,
}
_UZEX = {
	"custom_uzex_lot_no": "26111006497638",
	"custom_uzex_customer_org": "Maktab #1",
	"custom_uzex_deadline": "2026-07-08 11:47:09",
	"custom_uzex_start_price": 6300000000.0,
	"custom_uzex_portal": "etender",
}
_COMPANY = {"name": "Anjan Surprise LLC", "tax_id": "301234567", "address": "Toshkent"}


class TestFormatters(unittest.TestCase):
	def test_fmt_date_ddmmyyyy(self):
		self.assertEqual(fmt_date("2026-07-08 11:47:09"), "08.07.2026")
		self.assertEqual(fmt_date("2026-07-08"), "08.07.2026")
		self.assertEqual(fmt_date(None), "—")

	def test_fmt_money_space_thousands(self):
		self.assertEqual(fmt_money(6300000000), "6 300 000 000")
		self.assertEqual(fmt_money(None), "0")  # 0 -> "0"


class TestAssemble(unittest.TestCase):
	def test_ready_when_complete(self):
		pkg = assemble_bid_package("DEAL-1", _INTAKE, _PNL, _UZEX, _COMPANY, "UZS")
		self.assertTrue(pkg["ready"])
		self.assertEqual(pkg["missing"], [])
		self.assertEqual(pkg["lot"]["lot_no"], "26111006497638")
		self.assertEqual(pkg["pricing"]["ostatok"], 500000000.0)
		self.assertEqual(pkg["company"]["name"], "Anjan Surprise LLC")

	def test_missing_lists_gaps(self):
		pkg = assemble_bid_package("DEAL-2", {}, {"bid_price": 0}, {}, {}, "UZS")
		for gap in ("Lot no", "Buyer", "Bid deadline", "Bid price", "Company name"):
			self.assertIn(gap, pkg["missing"])
		self.assertFalse(pkg["ready"])

	def test_uzex_fallback_when_intake_empty(self):
		pkg = assemble_bid_package("DEAL-3", {}, _PNL, _UZEX, _COMPANY, "UZS")
		# lot_no/buyer/deadline fall back to the uzex fields
		self.assertEqual(pkg["lot"]["lot_no"], "26111006497638")
		self.assertEqual(pkg["lot"]["buyer"], "Maktab #1")
		self.assertTrue(pkg["ready"])


try:
	import docx  # noqa: F401

	_HAS_DOCX = True
except ImportError:
	_HAS_DOCX = False


@unittest.skipUnless(_HAS_DOCX, "python-docx not installed")
class TestBuildDocx(unittest.TestCase):
	def test_docx_opens_and_has_fields(self):
		from docx import Document

		from stabler.api._bid_package import build_bid_docx

		pkg = assemble_bid_package("DEAL-1", _INTAKE, _PNL, _UZEX, _COMPANY, "UZS")
		with tempfile.TemporaryDirectory() as d:
			path = os.path.join(d, "bid.docx")
			build_bid_docx(pkg, path, today=date(2026, 7, 8))
			self.assertTrue(os.path.exists(path) and os.path.getsize(path) > 0)
			# Reopen and confirm the merged text carries the key fields.
			doc = Document(path)
			text = "\n".join(p.text for p in doc.paragraphs)
			for t in doc.tables:
				for row in t.rows:
					text += "\n" + " | ".join(c.text for c in row.cells)
			self.assertIn("26111006497638", text)
			self.assertIn("Anjan Surprise LLC", text)
			self.assertIn("08.07.2026", text)  # dd.mm.yyyy
			self.assertIn("6 300 000 000", text)  # space thousands
			self.assertIn("Остаток", text)


if __name__ == "__main__":
	unittest.main()
