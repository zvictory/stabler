"""Bid-package assembly for UZEX tenders (WP-306, Frappe-free core).

Collects everything a human needs to submit a tender bid into one JSON:
lot info, the priced bid + margin/Остаток breakdown, deadline, company details,
and a ``missing[]`` list of gaps so an incomplete package is visible BEFORE it
goes out. The Frappe endpoint (api/tender.py::bid_package) hydrates the inputs
and renders the docx; this module only shapes + validates, so it unit-tests
without a database.

Deliberate scope limit: this prepares a package for a HUMAN to sign (E-IMZO) and
upload on the portal. There is no automated submission — see
docs/plans/uzex-eimzo-feasibility.md.
"""

from __future__ import annotations

from typing import Any


def _first(*vals):
	for v in vals:
		if v not in (None, "", 0, 0.0):
			return v
	return None


def fmt_date(value) -> str:
	"""Render a stored date/datetime as dd.mm.yyyy (Stabler date standard)."""
	if not value:
		return "—"
	s = str(value).strip().replace("T", " ")
	d = s.split(" ")[0]  # date part
	parts = d.split("-")
	if len(parts) == 3 and len(parts[0]) == 4:
		return f"{parts[2]}.{parts[1]}.{parts[0]}"
	return d or "—"


def fmt_money(value) -> str:
	"""Whole-number amount with space thousands (e.g. 162 975 654)."""
	try:
		n = float(value or 0)
	except (TypeError, ValueError):
		return "—"
	return f"{n:,.0f}".replace(",", " ")


def assemble_bid_package(
	deal: str,
	intake: dict | None,
	pnl: dict | None,
	uzex: dict | None,
	company: dict | None,
	currency: str = "",
) -> dict:
	"""Shape the submission dataset and compute ``missing[]``.

	Args mirror what the endpoint pulls: ``intake`` = custom_tender_intake JSON,
	``pnl`` = _compute_bid_pnl output, ``uzex`` = the custom_uzex_* fields,
	``company`` = {name, tax_id, address}.
	"""
	intake = intake or {}
	pnl = pnl or {}
	uzex = uzex or {}
	company = company or {}

	lot = {
		"lot_no": _first(uzex.get("custom_uzex_lot_no"), intake.get("lot_no")),
		"buyer": _first(intake.get("buyer"), uzex.get("custom_uzex_customer_org")),
		"bid_deadline": _first(intake.get("bid_deadline"), uzex.get("custom_uzex_deadline")),
		"start_price": _first(uzex.get("custom_uzex_start_price")),
		"portal": uzex.get("custom_uzex_portal"),
	}
	bid_price = pnl.get("bid_price")
	pricing = {
		"bid_price": bid_price,
		"currency": currency,
		"net_revenue": pnl.get("net_revenue"),
		"vat": pnl.get("vat"),
		"landed_goods": pnl.get("landed_goods"),
		"above_total": pnl.get("above_total"),
		"profit": pnl.get("profit"),
		"ostatok": pnl.get("ostatok"),
		"margin_on_revenue_pct": pnl.get("margin_on_revenue_pct"),
	}
	comp = {
		"name": company.get("name") or company.get("company_name"),
		"tax_id": company.get("tax_id"),
		"address": company.get("address"),
	}

	# Required fields for a submittable package → human-readable gaps.
	missing: list[str] = []
	if not lot["lot_no"]:
		missing.append("Lot no")
	if not lot["buyer"]:
		missing.append("Buyer")
	if not lot["bid_deadline"]:
		missing.append("Bid deadline")
	if not bid_price or float(bid_price or 0) <= 0:
		missing.append("Bid price")
	if not comp["name"]:
		missing.append("Company name")

	return {
		"deal": deal,
		"lot": lot,
		"pricing": pricing,
		"company": comp,
		"missing": missing,
		"ready": not missing,
	}


def build_bid_docx(package: dict, path: str, today=None) -> str:
	"""Render the bid letter + price-offer table to a .docx at ``path``.

	python-docx is imported lazily so the pure assembly above unit-tests without
	the dependency. Dates render dd.mm.yyyy. Returns ``path``.
	"""
	from datetime import date

	from docx import Document  # lazy: prod bench must have python-docx

	today = today or date.today()
	lot = package.get("lot") or {}
	pricing = package.get("pricing") or {}
	company = package.get("company") or {}
	ccy = pricing.get("currency") or ""

	doc = Document()
	doc.add_heading("Tender taklifi", level=0)
	doc.add_paragraph(f"Sana: {today.strftime('%d.%m.%Y')}")
	comp_line = company.get("name") or "—"
	if company.get("tax_id"):
		comp_line += f" (STIR: {company['tax_id']})"
	doc.add_paragraph(comp_line)
	if company.get("address"):
		doc.add_paragraph(str(company["address"]))

	doc.add_heading("Lot ma'lumoti", level=1)
	lot_rows = [
		("Lot №", lot.get("lot_no") or "—"),
		("Buyurtmachi", lot.get("buyer") or "—"),
		("Taklif muddati", fmt_date(lot.get("bid_deadline"))),
		("Boshlang'ich narx", f"{fmt_money(lot.get('start_price'))} {ccy}".strip()),
	]
	_two_col_table(doc, lot_rows)

	doc.add_heading("Narx taklifi", level=1)
	price_rows = [
		("Taklif narxi (Договор)", f"{fmt_money(pricing.get('bid_price'))} {ccy}".strip()),
		("НДС (QQS)", f"{fmt_money(pricing.get('vat'))} {ccy}".strip()),
		("Sof daromad", f"{fmt_money(pricing.get('net_revenue'))} {ccy}".strip()),
		("Tovar tannarxi (landed)", f"{fmt_money(pricing.get('landed_goods'))} {ccy}".strip()),
		("Foyda", f"{fmt_money(pricing.get('profit'))} {ccy}".strip()),
		("Остаток", f"{fmt_money(pricing.get('ostatok'))} {ccy}".strip()),
		("Marja (daromadga, %)", f"{pricing.get('margin_on_revenue_pct') if pricing.get('margin_on_revenue_pct') is not None else '—'}"),
	]
	_two_col_table(doc, price_rows)

	doc.add_paragraph("")
	note = doc.add_paragraph(
		"Ushbu paket imzolash uchun tayyorlangan. E-IMZO kaliti bilan imzolab, "
		"UZEX portaliga qo'lda yuklang (avtomatik yuborilmaydi)."
	)
	note.runs[0].italic = True
	doc.save(path)
	return path


def _two_col_table(doc, rows) -> None:
	table = doc.add_table(rows=len(rows), cols=2)
	table.style = "Light Grid Accent 1"
	for i, (label, value) in enumerate(rows):
		table.rows[i].cells[0].text = str(label)
		table.rows[i].cells[1].text = str(value)
